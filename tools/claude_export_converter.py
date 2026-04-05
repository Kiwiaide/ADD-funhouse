#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Claude 对话导出转换器 v3
conversations.json → Markdown / Word / PDF
"""

import json, os, sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from datetime import datetime

# Windows 高分屏 DPI 修复（必须在创建窗口前调用）
if sys.platform == "win32":
    try:
        import ctypes
        ctypes.windll.shcore.SetProcessDpiAwareness(2)  # Per-Monitor DPI aware
    except Exception:
        try:
            ctypes.windll.user32.SetProcessDPIAware()
        except Exception:
            pass

# ── 可选依赖 ──────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# ── 配色 ──────────────────────────────────────────────────────────
BG       = "#FAF9F5"
BG2      = "#F0EEE8"
FG       = "#1A1A1A"
FG2      = "#6B6B6B"
ACCENT   = "#D97706"
ACCENT_H = "#B45309"
BTN_FG   = "#FFFFFF"
BORDER   = "#DDD9D0"
SEL_BG   = "#FDE68A"
FONT_B   = ("Segoe UI", 10) if sys.platform == "win32" else ("SF Pro Text", 10)
FONT_H   = ("Segoe UI Semibold", 10) if sys.platform == "win32" else ("SF Pro Text", 10)
FONT_BIG = ("Segoe UI Semibold", 12) if sys.platform == "win32" else ("SF Pro Text", 12)


# ── 工具函数 ──────────────────────────────────────────────────────
def fmt_time(s):
    try:
        dt = datetime.fromisoformat(s.replace("Z", "+00:00"))
        return dt.astimezone().strftime("%Y-%m-%d %H:%M")
    except Exception:
        return s or ""


def safe_fn(name, n=60):
    if not name:
        name = "未命名对话"
    bad = set(r'\/:*?"<>|')
    r = "".join("_" if c in bad else c for c in name)[:n].strip("_. ")
    return r or "未命名对话"


def extract_blocks(msg):
    """从消息中提取 thinking 和 text 块，返回 [(type, content), ...]"""
    content = msg.get("content")
    if isinstance(content, list):
        blocks = []
        for block in content:
            if not isinstance(block, dict):
                continue
            btype = block.get("type", "")
            if btype == "thinking":
                text = (block.get("thinking") or "").strip()
                if text:
                    blocks.append(("thinking", text))
            elif btype == "text":
                text = (block.get("text") or "").strip()
                if text:
                    blocks.append(("text", text))
        if blocks:
            return blocks
    # fallback: 没有 content 数组时用 text 字段
    text = (msg.get("text") or "").strip()
    if text:
        return [("text", text)]
    return []


def load_conversations(fp):
    with open(fp, "r", encoding="utf-8") as f:
        data = json.load(f)
    return [c for c in data if c.get("chat_messages")]


# ── HTML ─────────────────────────────────────────────────────────
_HTML_STYLE = """
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    font-family: "Segoe UI", "PingFang SC", "Microsoft YaHei", sans-serif;
    background: #FAF9F5; color: #1A1A1A;
    max-width: 800px; margin: 0 auto; padding: 32px 24px;
  }
  h1 { font-size: 1.5em; font-weight: 600; margin-bottom: 4px; }
  .meta { color: #999; font-size: 0.85em; margin-bottom: 28px; }
  .msg { border-radius: 10px; padding: 12px 16px; margin-bottom: 10px; }
  .msg-label {
    font-size: 0.78em; font-weight: 600; margin-bottom: 6px;
    letter-spacing: 0.03em;
  }
  .msg-time { font-weight: 400; color: #999; margin-left: 8px; }
  .msg-human { background: #E8F0FA; }
  .msg-human .msg-label { color: #21448A; }
  .msg-claude { background: #FFF3E0; }
  .msg-claude .msg-label { color: #B45309; }
  .msg p { font-size: 0.96em; line-height: 1.65; white-space: pre-wrap; }
  .thinking { margin: 8px 0; border-left: 3px solid #C4B5A0; padding: 8px 12px;
    background: rgba(0,0,0,0.03); border-radius: 4px; }
  .thinking summary { cursor: pointer; color: #8B7355; font-size: 0.85em;
    font-weight: 600; margin-bottom: 4px; }
  .thinking p { font-size: 0.88em; color: #6B6B6B; line-height: 1.55; }
  hr { border: none; border-top: 1px solid #DDD9D0; margin: 28px 0; }
  @media print {
    body { padding: 0; }
    .msg { break-inside: avoid; }
  }
</style>
"""

def conv_to_html(conv):
    import html as html_lib
    title   = html_lib.escape(conv.get("name") or "未命名对话")
    created = fmt_time(conv.get("created_at", ""))
    parts   = [f"<!DOCTYPE html><html><head><meta charset='utf-8'>",
               f"<title>{title}</title>{_HTML_STYLE}</head><body>",
               f"<h1>{title}</h1>",
               f"<p class='meta'>创建时间：{created}</p>"]

    for msg in conv.get("chat_messages", []):
        blocks = extract_blocks(msg)
        if not blocks:
            continue
        is_human  = msg.get("sender") == "human"
        cls       = "msg-human" if is_human else "msg-claude"
        label     = "用户" if is_human else "Claude"
        ts        = html_lib.escape(fmt_time(msg.get("created_at", "")))
        parts.append(f"<div class='msg {cls}'>")
        parts.append(f"  <div class='msg-label'>{label}<span class='msg-time'>{ts}</span></div>")
        for btype, text in blocks:
            safe_text = html_lib.escape(text)
            if btype == "thinking":
                parts.append(f"  <details class='thinking'>")
                parts.append(f"    <summary>💭 思考过程</summary>")
                parts.append(f"    <p>{safe_text}</p>")
                parts.append(f"  </details>")
            else:
                parts.append(f"  <p>{safe_text}</p>")
        parts.append("</div>")

    parts.append("</body></html>")
    return "\n".join(parts)


# ── Markdown ──────────────────────────────────────────────────────
def conv_to_md(conv):
    title   = conv.get("name") or "未命名对话"
    created = fmt_time(conv.get("created_at", ""))
    lines   = [f"# {title}", "", f"> 创建时间：{created}", ""]
    for msg in conv.get("chat_messages", []):
        blocks = extract_blocks(msg)
        if not blocks:
            continue
        is_human = msg.get("sender") == "human"
        ts = fmt_time(msg.get("created_at", ""))
        label = f"**用户** ({ts})" if is_human else f"**Claude** ({ts})"
        lines += ["---", label, ""]
        for btype, text in blocks:
            if btype == "thinking":
                lines += ["<details>", "<summary>💭 思考过程</summary>", "",
                           text, "", "</details>", ""]
            else:
                lines += [text, ""]
    return "\n".join(lines)


# ── Word 工具 ─────────────────────────────────────────────────────
def _cjk(run, fn="微软雅黑"):
    rPr   = run._r.get_or_add_rPr()
    fonts = rPr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rPr.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), fn)
    fonts.set(qn("w:ascii"),    "Calibri")
    fonts.set(qn("w:hAnsi"),    "Calibri")


def _shade(cell, hex_color):
    """给表格单元格加背景色"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    """设置单元格边框（用于去除边框）"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:val"),   kwargs.get(side, "nil"))
        node.set(qn("w:sz"),    "0")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "auto")
        tcBorders.append(node)
    tcPr.append(tcBorders)


def _add_para(cell_or_doc, text, bold=False, size=10.5,
              color=None, space_after=0, italic=False):
    p   = cell_or_doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after  = Pt(space_after)
    fmt.line_spacing = Pt(size * 1.5)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        if DOCX_AVAILABLE:
            _cjk(run)
    return p


def conv_to_docx(conv):
    doc = Document()

    # 全局字体
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 页边距
    section = doc.sections[0]
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    title   = conv.get("name") or "未命名对话"
    created = fmt_time(conv.get("created_at", ""))

    # 标题
    h = doc.add_heading("", level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after  = Pt(4)
    run = h.add_run(title)
    run.font.size = Pt(15)
    _cjk(run)

    # 元信息
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after  = Pt(12)
    run = meta.add_run(f"创建时间：{created}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    _cjk(run)

    # 每条消息用单列表格实现气泡感
    for msg in conv.get("chat_messages", []):
        blocks = extract_blocks(msg)
        if not blocks:
            continue
        is_human = msg.get("sender") == "human"

        # 表格：1行1列
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.cell(0, 0)

        # 背景色：用户=浅蓝，Claude=浅橙
        bg_hex = "E8F0FA" if is_human else "FFF3E0"
        _shade(cell, bg_hex)
        _set_cell_border(cell)

        # 角色标签
        label_color = RGBColor(0x21, 0x6E, 0xB4) if is_human \
            else RGBColor(0xB4, 0x53, 0x09)
        label_text  = "用户" if is_human else "Claude"

        ts = fmt_time(msg.get("created_at", ""))
        lp = cell.add_paragraph()
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(2)
        lr = lp.add_run(label_text)
        lr.bold = True
        lr.font.size  = Pt(9)
        lr.font.color.rgb = label_color
        _cjk(lr)
        if ts:
            tr = lp.add_run(f"  {ts}")
            tr.font.size = Pt(8)
            tr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            _cjk(tr)

        # 正文（区分 thinking 和 text）
        for btype, text in blocks:
            if btype == "thinking":
                # 思考过程标签
                tp = cell.add_paragraph()
                tp.paragraph_format.space_before = Pt(4)
                tp.paragraph_format.space_after  = Pt(2)
                tr = tp.add_run("💭 思考过程")
                tr.bold = True
                tr.italic = True
                tr.font.size = Pt(8.5)
                tr.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
                _cjk(tr)
                for line in text.split("\n"):
                    lp2 = cell.add_paragraph(line)
                    lp2.paragraph_format.space_before = Pt(0)
                    lp2.paragraph_format.space_after  = Pt(0)
                    lp2.paragraph_format.line_spacing = Pt(14)
                    if lp2.runs:
                        lp2.runs[0].font.size = Pt(9)
                        lp2.runs[0].font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)
                        lp2.runs[0].italic = True
                        _cjk(lp2.runs[0])
                # 分隔线
                sep = cell.add_paragraph()
                sep.paragraph_format.space_before = Pt(2)
                sep.paragraph_format.space_after  = Pt(2)
                sr = sep.add_run("── 回复 ──")
                sr.font.size = Pt(8)
                sr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
                _cjk(sr)
            else:
                for line in text.split("\n"):
                    lp2 = cell.add_paragraph(line)
                    lp2.paragraph_format.space_before = Pt(0)
                    lp2.paragraph_format.space_after  = Pt(0)
                    lp2.paragraph_format.line_spacing = Pt(15)
                    if lp2.runs:
                        lp2.runs[0].font.size = Pt(10.5)
                        _cjk(lp2.runs[0])

        ep = cell.add_paragraph()
        ep.paragraph_format.space_before = Pt(0)
        ep.paragraph_format.space_after  = Pt(4)

        # 表格后加小间距段落
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(4)

    return doc


# ── PDF ───────────────────────────────────────────────────────────
def _get_cjk_font():
    """尝试注册系统中文字体，返回字体名称"""
    paths = [
        r"C:\Windows\Fonts\msyh.ttc",          # 微软雅黑
        r"C:\Windows\Fonts\simhei.ttf",         # 黑体
        r"C:\Windows\Fonts\simsun.ttc",         # 宋体
        "/System/Library/Fonts/PingFang.ttc",   # Mac PingFang
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",  # Linux
    ]
    for p in paths:
        if os.path.exists(p):
            try:
                name = "CJKFont"
                pdfmetrics.registerFont(TTFont(name, p))
                return name
            except Exception:
                continue
    return None


def conv_to_pdf(conv, path):
    if not PDF_AVAILABLE:
        raise ImportError("需要安装 reportlab：pip install reportlab")

    cjk = _get_cjk_font()
    font = cjk if cjk else "Helvetica"

    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm
    )

    USER_BG   = colors.HexColor("#E8F0FA")
    CLAUDE_BG = colors.HexColor("#FFF3E0")
    USER_FG   = colors.HexColor("#21448A")
    CLAUDE_FG = colors.HexColor("#B45309")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title2",
        fontName=font, fontSize=15, leading=20,
        textColor=colors.HexColor("#1A1A1A"),
        spaceAfter=4)
    meta_style = ParagraphStyle("Meta",
        fontName=font, fontSize=9, leading=12,
        textColor=colors.HexColor("#999999"),
        spaceAfter=14)
    label_h_style = ParagraphStyle("LabelH",
        fontName=font, fontSize=9, leading=13,
        textColor=USER_FG, spaceAfter=2,
        backColor=USER_BG, borderPadding=(4, 6, 0, 6))
    label_c_style = ParagraphStyle("LabelC",
        fontName=font, fontSize=9, leading=13,
        textColor=CLAUDE_FG, spaceAfter=2,
        backColor=CLAUDE_BG, borderPadding=(4, 6, 0, 6))
    body_h_style = ParagraphStyle("BodyH",
        fontName=font, fontSize=10.5, leading=16,
        textColor=colors.HexColor("#1A1A1A"),
        backColor=USER_BG, borderPadding=(0, 6, 6, 6))
    body_c_style = ParagraphStyle("BodyC",
        fontName=font, fontSize=10.5, leading=16,
        textColor=colors.HexColor("#1A1A1A"),
        backColor=CLAUDE_BG, borderPadding=(0, 6, 6, 6))

    story = []
    title   = conv.get("name") or "未命名对话"
    created = fmt_time(conv.get("created_at", ""))
    story.append(Paragraph(title,              title_style))
    story.append(Paragraph(f"创建时间：{created}", meta_style))

    think_label_style = ParagraphStyle("ThinkLabel",
        fontName=font, fontSize=8.5, leading=12,
        textColor=colors.HexColor("#8B7355"), spaceAfter=2,
        backColor=colors.HexColor("#FFF3E0"), borderPadding=(4, 6, 0, 6))
    think_body_style = ParagraphStyle("ThinkBody",
        fontName=font, fontSize=9, leading=13,
        textColor=colors.HexColor("#6B6B6B"),
        backColor=colors.HexColor("#FFF8F0"), borderPadding=(0, 6, 2, 6))
    separator_style = ParagraphStyle("Sep",
        fontName=font, fontSize=8, leading=10,
        textColor=colors.HexColor("#AAAAAA"),
        backColor=colors.HexColor("#FFF3E0"), borderPadding=(2, 6, 2, 6),
        alignment=1)

    for msg in conv.get("chat_messages", []):
        blocks = extract_blocks(msg)
        if not blocks:
            continue
        is_human = msg.get("sender") == "human"
        label    = "用户" if is_human else "Claude"

        ts = fmt_time(msg.get("created_at", ""))
        ts_suffix = f"  <font size='7' color='#999999'>{ts}</font>" if ts else ""

        if is_human:
            story.append(Paragraph(f"{label}{ts_suffix}", label_h_style))
            for btype, text in blocks:
                for line in text.split("\n"):
                    safe = line.replace("&", "&amp;").replace("<", "&lt;")
                    story.append(Paragraph(safe or " ", body_h_style))
        else:
            story.append(Paragraph(f"{label}{ts_suffix}", label_c_style))
            for btype, text in blocks:
                if btype == "thinking":
                    story.append(Paragraph("💭 思考过程", think_label_style))
                    for line in text.split("\n"):
                        safe = line.replace("&", "&amp;").replace("<", "&lt;")
                        story.append(Paragraph(safe or " ", think_body_style))
                    story.append(Paragraph("── 回复 ──", separator_style))
                else:
                    for line in text.split("\n"):
                        safe = line.replace("&", "&amp;").replace("<", "&lt;")
                        story.append(Paragraph(safe or " ", body_c_style))
        story.append(Spacer(1, 0.25*cm))

    doc.build(story)


# ── 保存控制器 ────────────────────────────────────────────────────
def save_convs(convs, output_dir, fmt, merge, merge_name="合并导出"):
    saved = []

    if merge:
        name = safe_fn(merge_name)
        if fmt == "html":
            parts = []
            for i, conv in enumerate(convs):
                # 取 body 内容
                body = conv_to_html(conv).split("<body>", 1)[1].rsplit("</body>", 1)[0]
                if i > 0:
                    parts.append("<hr>")
                parts.append(body)
            import html as html_lib
            content = (f"<!DOCTYPE html><html><head><meta charset='utf-8'>"
                       f"<title>{html_lib.escape(merge_name)}</title>"
                       f"{_HTML_STYLE}</head><body>" + "\n".join(parts) + "</body></html>")
            path = os.path.join(output_dir, f"{name}.html")
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            saved.append(path)

        elif fmt == "md":
            path = os.path.join(output_dir, f"{name}.md")
            with open(path, "w", encoding="utf-8") as f:
                f.write("\n\n---\n\n".join(conv_to_md(c) for c in convs))
            saved.append(path)

        elif fmt == "docx":
            if not DOCX_AVAILABLE:
                return None, "需安装 python-docx：pip install python-docx"
            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            for i, conv in enumerate(convs):
                if i > 0:
                    doc.add_page_break()
                for el in conv_to_docx(conv).element.body:
                    doc.element.body.append(el)
            path = os.path.join(output_dir, f"{name}.docx")
            doc.save(path)
            saved.append(path)

        elif fmt == "pdf":
            if not PDF_AVAILABLE:
                return None, "需安装 reportlab：pip install reportlab"
            # 多对话合并 PDF：逐个生成再合并
            try:
                from pypdf import PdfMerger
                merger = PdfMerger()
                tmp_files = []
                for i, conv in enumerate(convs):
                    tmp = os.path.join(output_dir, f"_tmp_{i}.pdf")
                    conv_to_pdf(conv, tmp)
                    merger.append(tmp)
                    tmp_files.append(tmp)
                path = os.path.join(output_dir, f"{name}.pdf")
                merger.write(path)
                merger.close()
                for t in tmp_files:
                    os.remove(t)
            except ImportError:
                # 没有 pypdf，就生成多个
                for i, conv in enumerate(convs):
                    cn = safe_fn(conv.get("name") or f"对话{i+1}")
                    path = os.path.join(output_dir, f"{cn}.pdf")
                    conv_to_pdf(conv, path)
                    saved.append(path)
                return saved, None
            saved.append(path)

    else:
        for conv in convs:
            cn = safe_fn(conv.get("name") or "未命名对话")
            if fmt == "html":
                path = os.path.join(output_dir, f"{cn}.html")
                with open(path, "w", encoding="utf-8") as f:
                    f.write(conv_to_html(conv))
            elif fmt == "md":
                path = os.path.join(output_dir, f"{cn}.md")
                with open(path, "w", encoding="utf-8") as f:
                    f.write(conv_to_md(conv))
            elif fmt == "docx":
                if not DOCX_AVAILABLE:
                    return None, "需安装 python-docx：pip install python-docx"
                path = os.path.join(output_dir, f"{cn}.docx")
                conv_to_docx(conv).save(path)
            elif fmt == "pdf":
                if not PDF_AVAILABLE:
                    return None, "需安装 reportlab：pip install reportlab"
                path = os.path.join(output_dir, f"{cn}.pdf")
                conv_to_pdf(conv, path)
            saved.append(path)

    return saved, None


# ── GUI ───────────────────────────────────────────────────────────
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Claude 对话转换器")
        self.configure(bg=BG)
        self.resizable(True, True)
        self.minsize(640, 540)
        self.conversations = []
        self._style()
        self._build_ui()

    def _style(self):
        s = ttk.Style(self)
        s.theme_use("clam")
        for w in ("TFrame", "TLabel", "TLabelframe", "TLabelframe.Label",
                  "TRadiobutton", "TCheckbutton"):
            s.configure(w, background=BG, foreground=FG, font=FONT_B,
                        bordercolor=BORDER)
        s.configure("TEntry", fieldbackground=BG2, foreground=FG, font=FONT_B)
        s.configure("Accent.TButton", background=ACCENT, foreground=BTN_FG,
                    font=FONT_BIG, padding=(20, 8), borderwidth=0, relief="flat")
        s.map("Accent.TButton",
              background=[("active", ACCENT_H), ("pressed", ACCENT_H)])
        s.configure("Ghost.TButton", background=BG2, foreground=FG2,
                    font=FONT_B, padding=(8, 4), borderwidth=0, relief="flat")
        s.map("Ghost.TButton", background=[("active", BORDER)])

    def _build_ui(self):
        # Header
        hdr = tk.Frame(self, bg=BG2, padx=20, pady=14)
        hdr.pack(fill="x")
        tk.Label(hdr, text="Claude 对话转换器",
                 font=("Segoe UI Semibold", 14), bg=BG2, fg=FG).pack(side="left")
        tk.Label(hdr, text="JSON → Markdown / Word / PDF",
                 font=FONT_B, bg=BG2, fg=FG2).pack(side="left", padx=12)

        body = tk.Frame(self, bg=BG, padx=16, pady=12)
        body.pack(fill="both", expand=True)

        # 步骤 1
        self._label(body, "1  选择 conversations.json 文件")
        f1 = tk.Frame(body, bg=BG); f1.pack(fill="x", pady=(0, 12))
        self.file_var = tk.StringVar(value="未选择文件")
        tk.Label(f1, textvariable=self.file_var, bg=BG2, fg=FG2,
                 font=FONT_B, anchor="w", padx=8, pady=6,
                 width=52).pack(side="left")
        ttk.Button(f1, text="浏览…", style="Ghost.TButton",
                   command=self._pick_file).pack(side="left", padx=6)

        # 步骤 2
        self._label(body, "2  选择要转换的对话")
        bar = tk.Frame(body, bg=BG); bar.pack(fill="x", pady=(0, 4))
        ttk.Button(bar, text="全选",  style="Ghost.TButton",
                   command=self._select_all).pack(side="left")
        ttk.Button(bar, text="全不选", style="Ghost.TButton",
                   command=self._deselect_all).pack(side="left", padx=6)

        lf = tk.Frame(body, bg=BORDER, bd=1)
        lf.pack(fill="both", expand=True, pady=(0, 12))
        inner = tk.Frame(lf, bg=BG2); inner.pack(fill="both", expand=True, padx=1, pady=1)
        sb = tk.Scrollbar(inner); sb.pack(side="right", fill="y")
        self.listbox = tk.Listbox(
            inner, selectmode="multiple", height=9,
            yscrollcommand=sb.set, activestyle="none",
            bg=BG2, fg=FG, selectbackground=SEL_BG, selectforeground=FG,
            font=FONT_B, relief="flat", bd=0, highlightthickness=0)
        self.listbox.pack(fill="both", expand=True)
        sb.config(command=self.listbox.yview)

        # 步骤 3
        self._label(body, "3  输出设置")
        opt = tk.Frame(body, bg=BG); opt.pack(fill="x", pady=(0, 8))

        fmtf = tk.Frame(opt, bg=BG); fmtf.pack(anchor="w")
        tk.Label(fmtf, text="格式：", bg=BG, fg=FG2, font=FONT_B).pack(side="left")
        self.fmt_var = tk.StringVar(value="md")
        ttk.Radiobutton(fmtf, text="Markdown", variable=self.fmt_var,
                        value="md").pack(side="left")
        ttk.Radiobutton(fmtf,
                        text="Word" if DOCX_AVAILABLE else "Word ⚠需安装",
                        variable=self.fmt_var, value="docx").pack(side="left", padx=10)
        ttk.Radiobutton(fmtf,
                        text="PDF" if PDF_AVAILABLE else "PDF ⚠需安装",
                        variable=self.fmt_var, value="pdf").pack(side="left", padx=10)
        ttk.Radiobutton(fmtf, text="HTML（支持 emoji，可在浏览器打印成 PDF）",
                        variable=self.fmt_var, value="html").pack(side="left", padx=10)

        mf = tk.Frame(opt, bg=BG); mf.pack(anchor="w", pady=4)
        self.merge_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(mf, text="合并成一个文件",
                        variable=self.merge_var,
                        command=self._toggle_merge).pack(side="left")
        self.merge_name_var = tk.StringVar(value="合并导出")
        self.merge_entry = ttk.Entry(mf, textvariable=self.merge_name_var,
                                     width=22, state="disabled")
        self.merge_entry.pack(side="left", padx=8)

        # 步骤 4
        self._label(body, "4  选择保存位置")
        f4 = tk.Frame(body, bg=BG); f4.pack(fill="x", pady=(0, 4))
        self.out_var = tk.StringVar(value=os.path.expanduser("~/Desktop"))
        tk.Label(f4, textvariable=self.out_var, bg=BG2, fg=FG2,
                 font=FONT_B, anchor="w", padx=8, pady=6,
                 width=52).pack(side="left")
        ttk.Button(f4, text="浏览…", style="Ghost.TButton",
                   command=self._pick_output).pack(side="left", padx=6)

        # 按钮
        bf = tk.Frame(self, bg=BG, pady=12); bf.pack()
        ttk.Button(bf, text="开始转换 →", style="Accent.TButton",
                   command=self._convert).pack()

    def _label(self, parent, text):
        tk.Label(parent, text=text, bg=BG, fg=FG2, font=FONT_H
                 ).pack(anchor="w", pady=(4, 2))

    def _pick_file(self):
        path = filedialog.askopenfilename(
            title="选择 conversations.json",
            filetypes=[("JSON 文件", "*.json"), ("所有文件", "*.*")])
        if not path:
            return
        try:
            self.conversations = load_conversations(path)
        except Exception as e:
            messagebox.showerror("错误", f"无法读取文件：{e}")
            return
        self.file_var.set(os.path.basename(path))
        self.listbox.delete(0, "end")
        for conv in self.conversations:
            name    = conv.get("name") or "未命名对话"
            created = fmt_time(conv.get("created_at", ""))
            n       = len(conv.get("chat_messages", []))
            self.listbox.insert("end", f"  {name}   {created}   {n} 条")
        self._select_all()

    def _pick_output(self):
        p = filedialog.askdirectory(title="选择保存目录")
        if p: self.out_var.set(p)

    def _select_all(self):   self.listbox.select_set(0, "end")
    def _deselect_all(self): self.listbox.select_clear(0, "end")

    def _toggle_merge(self):
        self.merge_entry.config(
            state="normal" if self.merge_var.get() else "disabled")

    def _convert(self):
        idx = self.listbox.curselection()
        if not idx:
            messagebox.showwarning("提示", "请先选择至少一个对话。")
            return
        convs      = [self.conversations[i] for i in idx]
        fmt        = self.fmt_var.get()
        merge      = self.merge_var.get()
        merge_name = self.merge_name_var.get().strip() or "合并导出"
        out_dir    = self.out_var.get()

        if not os.path.isdir(out_dir):
            messagebox.showerror("错误", f"保存目录不存在：{out_dir}")
            return

        saved, err = save_convs(convs, out_dir, fmt, merge, merge_name)
        if err:
            messagebox.showerror("错误", err)
            return

        names = "\n".join(f"  · {os.path.basename(p)}" for p in saved)
        messagebox.showinfo("完成",
                            f"转换完成！生成了 {len(saved)} 个文件：\n\n{names}"
                            f"\n\n保存位置：{out_dir}")
        import subprocess, webbrowser
        if fmt == "html" and len(saved) == 1:
            webbrowser.open(f"file:///{saved[0].replace(os.sep, '/')}")
        elif sys.platform == "darwin":
            subprocess.Popen(["open", out_dir])
        elif sys.platform == "win32":
            os.startfile(out_dir)


if __name__ == "__main__":
    app = App()
    app.mainloop()
